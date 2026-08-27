import sys
from pathlib import Path

sys.stdout.reconfigure(encoding='utf-8')
docx_path = Path(r"e:\Projects\VectorWorkspace\Materials\Генератор Школьного Расписания в Казахстане.docx")

try:
    import docx
    doc = docx.Document(docx_path)
    for p in doc.paragraphs:
        if p.text.strip():
            print(p.text)
    for t in doc.tables:
        for row in t.rows:
            print(" | ".join(cell.text.strip().replace('\n', ' ') for cell in row.cells))
except Exception as e:
    print(f"Error reading docx with python-docx: {e}")
    # Fallback zip extract
    import zipfile, xml.etree.ElementTree as ET
    with zipfile.ZipFile(docx_path) as z:
        xml_content = z.read('word/document.xml')
        tree = ET.fromstring(xml_content)
        texts = tree.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t')
        print("".join([t.text for t in texts if t.text]))
